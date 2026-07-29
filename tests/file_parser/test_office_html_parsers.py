"""Office 和 HTML 解析器测试。"""

import io
import unittest

from agent_hub.file_parser.config import ParserConfig
from agent_hub.file_parser.parsers.html import HtmlParser


class TestHtmlParser(unittest.TestCase):
    """HTML 解析器测试"""

    def setUp(self):
        self.config = ParserConfig()
        self.parser = HtmlParser(self.config)

    def test_parse_simple_html(self):
        html = b"<html><body><p>Hello, World!</p></body></html>"
        result = self.parser.parse(html, "test.html")

        self.assertEqual(result.metadata["format"], "html")
        self.assertIn("Hello, World!", result.text)

    def test_parse_html_with_title(self):
        html = b"<html><head><title>Test Page</title></head><body>Content</body></html>"
        result = self.parser.parse(html, "page.html")

        self.assertEqual(result.metadata["title"], "Test Page")

    def test_strip_scripts(self):
        html = b"""
        <html>
        <body>
            <p>Safe content</p>
            <script>alert('XSS')</script>
            <p>More safe content</p>
        </body>
        </html>
        """
        result = self.parser.parse(html, "xss.html")

        self.assertIn("Safe content", result.text)
        self.assertNotIn("alert", result.text)
        self.assertNotIn("XSS", result.text)

    def test_strip_styles(self):
        html = b"""
        <html>
        <head><style>body { color: red; }</style></head>
        <body><p>Content</p></body>
        </html>
        """
        result = self.parser.parse(html, "styled.html")

        self.assertIn("Content", result.text)
        self.assertNotIn("color", result.text)

    def test_extract_sections_by_headings(self):
        html = b"""
        <html>
        <body>
            <h1>Main Title</h1>
            <p>Intro paragraph</p>
            <h2>Section 1</h2>
            <p>Content for section 1</p>
            <h2>Section 2</h2>
            <p>Content for section 2</p>
        </body>
        </html>
        """
        result = self.parser.parse(html, "sections.html")

        # 应该按标题切分成 chunks
        titles = [c.title for c in result.chunks]
        self.assertIn("Section 1", titles)
        self.assertIn("Section 2", titles)

    def test_parse_utf8_html(self):
        html = '<html><body><p>你好，世界！</p></body></html>'.encode("utf-8")
        result = self.parser.parse(html, "chinese.html")

        self.assertIn("你好，世界！", result.text)

    def test_parse_html_with_charset(self):
        html = b"""
        <html>
        <head><meta charset="utf-8"></head>
        <body><p>Content</p></body>
        </html>
        """
        result = self.parser.parse(html, "charset.html")

        self.assertIn("Content", result.text)

    def test_parse_html_without_body(self):
        html = b"<p>Just a paragraph</p>"
        result = self.parser.parse(html, "fragment.html")

        self.assertIn("Just a paragraph", result.text)

    def test_clean_multiple_whitespace(self):
        html = b"""
        <html>
        <body>
            <p>Line 1</p>


            <p>Line 2</p>
        </body>
        </html>
        """
        result = self.parser.parse(html, "whitespace.html")

        # 不应该有过多的空行
        self.assertNotIn("\n\n\n", result.text)


class TestDocxParser(unittest.TestCase):
    """DOCX 解析器测试（需要 python-docx）"""

    def setUp(self):
        self.config = ParserConfig()

    def test_parse_simple_docx(self):
        """测试解析简单 DOCX（需要创建有效的 DOCX 文件）"""
        try:
            from docx import Document as DocxDocument

            from agent_hub.file_parser.parsers.office import DocxParser
        except ImportError:
            self.skipTest("python-docx not installed")

        # 创建一个简单的 DOCX
        doc = DocxDocument()
        doc.add_paragraph("Hello, World!")
        doc.add_paragraph("Second paragraph.")

        buffer = io.BytesIO()
        doc.save(buffer)
        docx_content = buffer.getvalue()

        parser = DocxParser(self.config)
        result = parser.parse(docx_content, "test.docx")

        self.assertEqual(result.metadata["format"], "docx")
        self.assertIn("Hello, World!", result.text)
        self.assertIn("Second paragraph", result.text)

    def test_parse_docx_with_table(self):
        """测试解析包含表格的 DOCX"""
        try:
            from docx import Document as DocxDocument

            from agent_hub.file_parser.parsers.office import DocxParser
        except ImportError:
            self.skipTest("python-docx not installed")

        doc = DocxDocument()
        doc.add_paragraph("Document with table")

        # 添加表格
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(1, 0).text = "Value 1"
        table.cell(1, 1).text = "Value 2"

        buffer = io.BytesIO()
        doc.save(buffer)
        docx_content = buffer.getvalue()

        parser = DocxParser(self.config)
        result = parser.parse(docx_content, "table.docx")

        self.assertIn("Header 1", result.text)
        self.assertIn("Value 1", result.text)
        self.assertGreater(result.metadata["table_count"], 0)


class TestXlsxParser(unittest.TestCase):
    """XLSX 解析器测试（需要 openpyxl）"""

    def setUp(self):
        self.config = ParserConfig()

    def test_parse_simple_xlsx(self):
        """测试解析简单 XLSX"""
        try:
            from openpyxl import Workbook

            from agent_hub.file_parser.parsers.office import XlsxParser
        except ImportError:
            self.skipTest("openpyxl not installed")

        # 创建一个简单的 XLSX
        wb = Workbook()
        ws = wb.active
        ws["A1"] = "Name"
        ws["B1"] = "Value"
        ws["A2"] = "Item 1"
        ws["B2"] = 100

        buffer = io.BytesIO()
        wb.save(buffer)
        xlsx_content = buffer.getvalue()

        parser = XlsxParser(self.config)
        result = parser.parse(xlsx_content, "test.xlsx")

        self.assertEqual(result.metadata["format"], "xlsx")
        self.assertIn("Name", result.text)
        self.assertIn("Item 1", result.text)

    def test_parse_xlsx_multiple_sheets(self):
        """测试解析多 sheet 的 XLSX"""
        try:
            from openpyxl import Workbook

            from agent_hub.file_parser.parsers.office import XlsxParser
        except ImportError:
            self.skipTest("openpyxl not installed")

        wb = Workbook()
        ws1 = wb.active
        ws1.title = "Sheet1"
        ws1["A1"] = "Sheet 1 Content"

        ws2 = wb.create_sheet("Sheet2")
        ws2["A1"] = "Sheet 2 Content"

        buffer = io.BytesIO()
        wb.save(buffer)
        xlsx_content = buffer.getvalue()

        parser = XlsxParser(self.config)
        result = parser.parse(xlsx_content, "multi.xlsx")

        self.assertEqual(result.metadata["sheet_count"], 2)
        self.assertIn("Sheet 1 Content", result.text)
        self.assertIn("Sheet 2 Content", result.text)

        # 每个 sheet 应该是一个 chunk
        self.assertEqual(len(result.chunks), 2)


if __name__ == "__main__":
    unittest.main()
