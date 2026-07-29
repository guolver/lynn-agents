"""Office 文档解析器：DOCX、XLSX。"""

from __future__ import annotations

import io

from .base import BaseParser, ParsedChunk, ParseResult


class DocxParser(BaseParser):
    """DOCX 解析器，使用 python-docx"""

    def parse(self, file_bytes: bytes, filename: str) -> ParseResult:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))

        paragraphs: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    paragraphs.append(row_text)

        full_text = "\n\n".join(paragraphs)
        full_text = self._truncate_text(full_text)

        chunks = self._split_by_length(full_text, filename)

        return ParseResult(
            text=full_text,
            chunks=chunks,
            metadata={
                "format": "docx",
                "size_bytes": len(file_bytes),
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables),
            },
        )


class XlsxParser(BaseParser):
    """XLSX 解析器，使用 openpyxl"""

    def parse(self, file_bytes: bytes, filename: str) -> ParseResult:
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)

        sheets_text: list[str] = []
        chunks: list[ParsedChunk] = []
        total_rows = 0

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            rows: list[str] = []

            for row in sheet.iter_rows(values_only=True):
                # 过滤空行
                cell_values = [str(cell) if cell is not None else "" for cell in row]
                if any(v.strip() for v in cell_values):
                    rows.append(" | ".join(cell_values))
                    total_rows += 1

                # 限制行数
                if total_rows > 10000:
                    break

            if rows:
                sheet_text = f"=== Sheet: {sheet_name} ===\n" + "\n".join(rows)
                sheets_text.append(sheet_text)

                chunks.append(
                    ParsedChunk(
                        title=f"{filename} - {sheet_name}",
                        content="\n".join(rows),
                        metadata={"sheet": sheet_name, "row_count": len(rows)},
                    )
                )

            if len(chunks) >= self.config.max_chunks:
                break

        wb.close()

        full_text = "\n\n".join(sheets_text)
        full_text = self._truncate_text(full_text)

        return ParseResult(
            text=full_text,
            chunks=chunks,
            metadata={
                "format": "xlsx",
                "size_bytes": len(file_bytes),
                "sheet_count": len(wb.sheetnames),
                "total_rows": total_rows,
            },
        )
